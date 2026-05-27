#!/usr/bin/env python3
"""
将6张专利附图插入到docx文件中。
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_image_after_paragraph(doc, search_text, image_path, caption, max_search=200):
    """在包含search_text的段落后插入图片和标题。"""
    found = False
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # 找到目标段落，在其后插入图片
            new_para = doc.paragraphs[i]._element
            
            # 创建图片段落
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(image_path, width=Inches(5.5))
            
            # 将图片段落移到目标段落后
            new_para.addnext(img_para._element)
            
            # 创建标题段落
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.font.size = Pt(10)
            cap_run.font.bold = False
            cap_run.font.name = '宋体'
            
            # 将标题段落移到图片段落后
            img_para._element.addnext(cap_para._element)
            
            # 空行
            blank_para = doc.add_paragraph()
            img_para._element.addnext(blank_para._element)
            
            found = True
            print(f"  ✓ 插入 {caption} 到 '{search_text[:50]}...'")
            return True
    
    if not found:
        print(f"  ✗ 未找到段落: '{search_text[:50]}...'")
    return found

def main():
    docx_path = Path("docs/专利技术交底书_一种面向多任务的遥感时空嵌入底座构建方法.docx")
    figures_dir = Path("docs/patent_figures")
    
    if not docx_path.exists():
        print(f"错误: 找不到 {docx_path}")
        sys.exit(1)
    
    doc = Document(docx_path)
    
    # 图片插入配置: (搜索文本, 图片路径, 标题)
    insertions = [
        # 图1: 系统架构图 → 3.2 系统架构部分
        ("**损失层**包含重建损失、反坍缩损失、一致性损失和时序对比损失",
         figures_dir / "fig1_system_architecture.png",
         "图1  系统整体架构图"),
        
        # 图2: 数据流图 → 3.4 系统流程部分
        ("第四阶段为优化。各损失加权求和后通过反向传播更新模型参数。",
         figures_dir / "fig2_data_flow.png",
         "图2  多源数据到多任务复用的数据流程图"),
        
        # 图3: 格网划分 → 3.3 网格划分层部分
        ("网格单元作为数据组织的基本单位，确保了多源数据在输入模型前已在同一空间框架下对齐",
         figures_dir / "fig3_grid_division.png",
         "图3  离散格网划分示意图"),
        
        # 图4: 专用输入头 → 3.3 专用输入头部分
        ("输入头的设计保证了系统的可扩展性——新增数据源时，只需增加对应的专用输入头，无需修改下游模块。",
         figures_dir / "fig4_input_heads.png",
         "图4  专用输入头示意图"),
        
        # 图5: 双窗口时序对比 → 3.4.2 双窗口时序对比部分
        ("反斜对角对比损失：将批次中对角线位置的样本对视为负样本",
         figures_dir / "fig5_dual_window.png",
         "图5  双窗口时序对比示意图"),
        
        # 图6: 两阶段策略对比 → 3.4.1 预归一化反坍缩部分
        ("训练阶段跳过归一化操作，直接在原始幅度空间中计算上述全部反坍缩损失",
         figures_dir / "fig6_two_stage.png",
         "图6  预归一化两阶段策略对比图"),
    ]
    
    print("开始插入图片到专利文档...")
    success_count = 0
    for search_text, img_path, caption in insertions:
        if not img_path.exists():
            print(f"  ✗ 图片不存在: {img_path}")
            continue
        if insert_image_after_paragraph(doc, search_text, str(img_path), caption):
            success_count += 1
    
    # 保存
    output_path = docx_path
    doc.save(output_path)
    print(f"\n完成! 成功插入 {success_count}/{len(insertions)} 张图片。")
    print(f"文档已保存到: {output_path}")

if __name__ == "__main__":
    main()
